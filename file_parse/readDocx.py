import docx

# Load the existing .docx file
doc = docx.Document("ResumeC.docx")

# Create a new document
new_doc = docx.Document()

# Iterate through paragraphs and runs in the existing document
for paragraph in doc.paragraphs:
    new_paragraph = new_doc.add_paragraph()
    for run in paragraph.runs:
        # Extract text and formatting
        text = run.text
        formatting = {
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            # Add more formatting properties as needed
        }

        # Process the text (add "A" to the end of each word)
        words = text.split()
        modified_words = [word + "A" for word in words]
        modified_text = " ".join(modified_words)

        # Create a new run in the new paragraph with the modified text and formatting
        new_run = new_paragraph.add_run(modified_text)
        new_run.bold = formatting["bold"]
        new_run.italic = formatting["italic"]
        new_run.underline = formatting["underline"]
        # Add more formatting properties as needed

# Save the new document
new_doc.save("modified_document.docx")
